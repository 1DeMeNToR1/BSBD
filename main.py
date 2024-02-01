import re
import sys
import psycopg2
from PyQt5 import QtCore, QtGui, QtWidgets, Qt
from PyQt5.QtWidgets import QApplication, QMainWindow, \
    QTabWidget, QVBoxLayout, QWidget, QTableWidget, \
    QTableWidgetItem, QPushButton, QDialog, QLabel, \
    QLineEdit, QMessageBox, QListWidget
from PyQt5.QtCore import Qt, pyqtSignal
import traceback
from docx import Document


def log_uncaught_exceptions(ex_cls, ex, tb): #error catcher
    text = '{}: {}:\n'.format(ex_cls.__name__, ex)
    text += ''.join(traceback.format_tb(tb))
    print(text)
    QtWidgets.QMessageBox.critical(None, 'Error', text)
    sys.exit()
sys.excepthook = log_uncaught_exceptions

class LoginDialog(QDialog):
    login_signal = pyqtSignal(str, str)

    def __init__(self):
        super().__init__()

        self.setWindowTitle('Вход')
        self.setGeometry(100, 100, 300, 150)

        layout = QVBoxLayout()

        self.label_login = QLabel('Логин:')
        self.edit_login = QLineEdit(self)
        layout.addWidget(self.label_login)
        layout.addWidget(self.edit_login)

        self.label_password = QLabel('Пароль:')
        self.edit_password = QLineEdit(self)
        self.edit_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.label_password)
        layout.addWidget(self.edit_password)

        self.button_login = QPushButton('Войти', self)
        self.button_login.clicked.connect(self.on_login_button_click)
        layout.addWidget(self.button_login)

        self.setLayout(layout)

    def on_login_button_click(self):
        login = self.edit_login.text()
        password = self.edit_password.text()
        self.login_signal.emit(login, password)
        self.accept()


class SecondWindow(QDialog):
    data_signal = pyqtSignal(list)

    def __init__(self, field_labels, parent=None):
        super().__init__(parent)

        self.setWindowTitle('Внесение параметров')
        self.setGeometry(400, 400, 300, 200)

        layout = QVBoxLayout()

        self.field_widgets = []
        for label in field_labels:
            label_widget = QLabel(label + ':', self)
            edit_widget = QLineEdit(self)
            self.field_widgets.append(edit_widget)

            layout.addWidget(label_widget)
            layout.addWidget(edit_widget)

        self.button_ok = QPushButton('OK', self)
        self.button_ok.clicked.connect(self.on_ok_button_click)
        layout.addWidget(self.button_ok)

        self.setLayout(layout)

    def on_ok_button_click(self):
        data = [widget.text() for widget in self.field_widgets]
        self.data_signal.emit(data)
        #self.accept()

class ThirdWindow(QDialog):
    def __init__(self, data, parent=None):
        super().__init__(parent)

        self.setWindowTitle('Результат')
        self.setGeometry(500, 500, 500, 300)

        layout = QVBoxLayout()
        list_widget = QListWidget(self)
        list_widget.addItems([str(item) for item in data])

        layout.addWidget(list_widget)
        self.setLayout(layout)

class DatabaseInterface(QMainWindow):
    def __init__(self):
        super(DatabaseInterface, self).__init__()

        self.setWindowTitle("Автосервис")
        self.setGeometry(100, 100, 1200, 1200)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        self.tab_widget = QTabWidget(self.central_widget)

        self.initialize_ui()

    def initialize_ui(self):
        access_1 = ['public']
        access_3 = ['public', 'schema_for_admin']
        access_2 = ['public', 'schema_for_sadmin', 'schema_for_admin']
        login_dialog = LoginDialog()
        login_dialog.login_signal.connect(self.login_info)
        login_dialog.exec_()

        cursor = self.connection.cursor()
        cursor.execute(f"""SELECT rolname
                            FROM pg_user
                            LEFT JOIN pg_auth_members ON pg_user.usesysid = pg_auth_members.member
                            LEFT JOIN pg_roles ON pg_auth_members.roleid = pg_roles.oid
                            WHERE usename = current_user;""")
        results = cursor.fetchall()
        self.connection.commit()
        cursor.close()
        if str(results[0][0]) == 'Сотрудник':
            data = access_1
        elif str(results[0][0]) == "Админ" or results[0][0] == None:
            data = access_2
        else:
            data = access_3
        shemma = ', '.join(f"'{item}'" for item in data)
        print(shemma)
        print(results[0][0])

        buttons_query = f"""
                SELECT 
                    p.proname AS function_name, 
                    n.nspname AS schema_name,
                    pg_get_function_identity_arguments(p.oid) AS arguments,
                    pg_get_function_result(p.oid) AS result_type
                FROM 
                    pg_proc p
                    JOIN pg_namespace n ON p.pronamespace = n.oid
                WHERE
                    n.nspname IN ({shemma})
                ORDER BY 
                    schema_name, function_name;
                """
        try:
            cursor = self.connection.cursor()
            cursor.execute(buttons_query)
            results = cursor.fetchall()
            self.connection.commit()
            cursor.close()
            button_spacing = 0
            tab = QWidget()
            layout = QVBoxLayout(tab)

            for result in results:
                function_name, schema_name, arguments, result_type = result
                if "pgp" in function_name:
                    continue
                button = QPushButton(f"{function_name}", self)
                button.clicked.connect(lambda _, fname=function_name, args=arguments: self.on_function_button_click(fname, args))
                button.setStyleSheet("text-align: center;")
                layout.addWidget(button)
                layout.addSpacing(button_spacing)
            self.tab_widget.addTab(tab, "Кнопки")
        except Exception as e:
            print(e)
            error = str(e).split(":")
            error_message = f"Произошла ошибка при выполнении запроса1: {str(error[len(error) - 1])}"
            QMessageBox.critical(self, "Ошибка", error_message)
            print(error_message)
            self.connection.rollback()
        finally:
            cursor.close()

        tables = self.get_all_tables()
        for table in tables:
            self.create_tab(table)

        layout = QVBoxLayout(self.central_widget)
        layout.addWidget(self.tab_widget)

    def handle_third_window_data(self, data):
        third_window = ThirdWindow(data)
        third_window.exec_()



    def on_function_button_click(self, function_name, argumentss):
        arguments = argumentss.split('"')
        #print(arguments)
        if len(arguments) != 1:
            field_labels = []
            for i in range(1, len(arguments)):
                if 'character' in arguments[i] or 'varying' in arguments[i] or 'integer' in arguments[i]\
                        or 'numeric' in arguments[i] or 'date' in arguments[i] or 'text' in arguments[i]\
                        or 'timestamp' in arguments[i] or arguments[i] == '':
                    continue
                else:
                    field_labels.append(arguments[i])

            second_window = SecondWindow(field_labels)
            second_window.data_signal.connect(lambda data: self.handle_second_window_data(data, function_name))
            second_window.exec_()
        else:
            self.handle_second_window_data('', function_name)

    def get_to_file(self, function_name, orderdata):
        try:
            # Создаем новый документ
            new_document = Document()
            # Проходим по каждому заказу в списке
            for order_data in orderdata:
                # Генерируем текст договора на основе данных заказа
                contract_text = (
                    f"Договор по заказу #{orderdata[0][0]}\n\n"
                    f"Дата: {orderdata[0][1]}\n"
                    f"Статус заказа: {orderdata[0][2]}\n"
                    f"Клиент: {orderdata[0][3]} {orderdata[0][4]}\n"
                    f"Автомобиль: {orderdata[0][5]} {orderdata[0][6]}\n\n"
                    f"Детали заказа:{order_data[7]} {order_data[8]}\n"
                    f"Запчасть:{order_data[7]} {order_data[8]}\n"
                    f"Дата платежа: {order_data[10]}"
                )
                # Вставляем текст договора в новый документ

                new_document.add_paragraph(contract_text)
                new_document.add_paragraph(f"Итоговая сумма платежа: {order_data[9]}\n")
            # Сохраняем результаты в новый файл
            new_document.save("договор_по_заказам_результат.docx")
        except Exception as e:
            print(e)
            error = str(e).split(":")
            error_message = f"Произошла ошибка при выгрузки: {str(error[len(error) - 1])}"
            QMessageBox.critical(self, "Ошибка выгрузки", error_message)
            self.connection.rollback()

    def parse_tuple_string(self, s):
        if isinstance(s, (str, bytes)):
            match = re.match(r'\((.*?)\)', s.decode('utf-8') if isinstance(s, bytes) else s)
            if match:
                values = match.group(1).replace('"', '').split(',')
                parsed_values = [int(value) if value.isdigit() else value for value in values]
                return parsed_values
            else:
                print(s)
                return None
        else:
            return s
    def add_call(self, function_name, result_string):
        try:
            cursor = self.connection.cursor()
            cursor.execute(f'CALL {function_name}({result_string});')
            self.connection.commit()
        except Exception as e:
            print(e)
            error = str(e).split(":")
            error_message = f"Произошла ошибка2: {str(error[len(error) - 1])}"
            QMessageBox.critical(self, "Ошибка", error_message)
            self.connection.rollback()
        finally:
            if cursor:
                cursor.close()
    def handle_second_window_data(self, data, function_name):
        cursor = None
        try:
            result_string = ', '.join(f"'{item}'" for item in data)
            if "Добавить" in function_name:
                print(result_string)
                print(function_name)
                self.add_call(function_name, result_string)
            else:
                print(result_string)
                cursor = self.connection.cursor()
                #cursor.callproc(f'{function_name}({result_string})')
                cursor.execute(f'SELECT {function_name}({result_string});')
                output = cursor.fetchall()
                self.connection.commit()
                cursor.close()
                result_lists = [self.parse_tuple_string(item[0]) for item in output if self.parse_tuple_string(item[0]) is not None]

                if "Выгрузить" in function_name:
                    self.get_to_file(function_name, result_lists)
                else:
                    self.handle_third_window_data(result_lists)
        except Exception as e:
            print(e)
            error = str(e).split(":")
            error_message = f"Произошла ошибка3: {str(error[len(error) - 1])}"
            QMessageBox.critical(self, "Ошибка", error_message)
            self.connection.rollback()
        finally:
            if cursor:
                cursor.close()
    def login_info(self, login, password):
        try:
            self.connection = psycopg2.connect(
                database="autoservice",
                user=login,
                password=password,

                #user="89125553555",
                #password="mex1",

                #user="89161234567",
                #password="glmas1",

                #user="84959876543",
                #password="admin1",

                user="postgres",
                password="172839",
                host="192.168.0.9"
            )
        except Exception as e:
            print(e)
            error = str(e).split(":")
            error_message = f"Произошла ошибка4: {str(error[len(error)-1])}"
            QMessageBox.critical(self, "Ошибка", error_message)

    def create_tab(self, table_name):
        tab = QWidget()
        layout = QVBoxLayout(tab)

        data = self.fetch_data_from_table(table_name)
        headers = [desc[0] for desc in self.get_table_columns(table_name)]

        table_widget = QTableWidget(self)
        table_widget.setRowCount(len(data))
        table_widget.setColumnCount(len(headers))
        table_widget.setHorizontalHeaderLabels(headers)

        for i, row_data in enumerate(data):
            for j, item_data in enumerate(row_data):
                item = QTableWidgetItem(str(item_data))
                item.setFlags(item.flags() | Qt.ItemIsEditable)
                table_widget.setItem(i, j, item)

        layout.addWidget(table_widget)
        self.tab_widget.addTab(tab, table_name)

    def fetch_data_from_table(self, table_name):
        cursor = self.connection.cursor()
        cursor.execute(f"SELECT * FROM {table_name};")
        data = cursor.fetchall()
        self.connection.commit()
        cursor.close()
        return data

    def get_table_columns(self, table_name):
        cursor = self.connection.cursor()
        cursor.execute(f"SELECT column_name FROM information_schema.columns WHERE table_name = '{table_name}';")
        columns = cursor.fetchall()
        self.connection.commit()
        cursor.close()
        return columns

    def get_all_tables(self):
        cursor = self.connection.cursor()
        cursor.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = 'public';")
        tables = [table[0] for table in cursor.fetchall()]
        self.connection.commit()
        cursor.close()
        return tables

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = DatabaseInterface()
    window.show()
    sys.exit(app.exec_())